# Libraries
import os
import json
import time
import requests
import pandas as pd
import re
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from reportlab.platypus import Table, TableStyle
from reportlab.lib import colors
from docx import Document
from requests.exceptions import HTTPError
from io import BytesIO,StringIO
from bs4 import BeautifulSoup
import requests
from langchain_groq import ChatGroq
from langchain_core.prompts import PromptTemplate
from pathlib import Path
current_dir = Path(__file__).parent

model_name = 'llama-3.1-70b-versatile'

# Entire syllabus generator pipeline with all functions in this class
class Syllabus_generator :

    def __init__(self,grade,subject,syllabus_type,instructions,path=""):
        self.grade = grade
        self.subject = subject
        self.Syllabus_type = syllabus_type
        self.instructions = instructions
        self.model = ChatGroq(model=model_name, temperature=0.3, api_key="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")
        self.path = path

    def read_text_file(self,filepath):

        with open(f"{self.path}{filepath}", 'r') as file:
            return file.read()

    def build_prompt(self,filepath):
        # build invididual promps for each sub part of syllabus
        template = self.read_text_file(filepath)

        prompt = PromptTemplate.from_template(template)
        return prompt

    def Validator(self,response):
        data = []
        try:
            data = json.loads(response)
        except json.JSONDecodeError as e:


            print("JSON Decode Error , Trying to correct the JSON")
            try:
                corrected_result = response[min(response.find('{'),response.find('[')):max(response.rfind(']'),response.rfind('}')) + 1]
                data = json.loads(corrected_result)
                print("Corrected Parsed JSON successfully")
            except json.JSONDecodeError as e:
                print("Failed to parse corrected JSON")
        return data

    def course_description(self) -> str:
        # Brief description or high level overview of course

        # Build prompt
        prompt = self.build_prompt(current_dir / 'prompts/course_description.txt')
        # Create LLM chain
        chain = prompt | self.model
        # Generate response
        response = chain.invoke({"grade" : self.grade,
                                 "subject" : self.subject,
                                 "Syllabus_type" : self.Syllabus_type,
                                 "instructions" : self.instructions
                                 }).content
        return response

    def course_objectives(self,course_description:str) -> str:
        # Numbered objectives or aims of the course

        prompt = self.build_prompt(current_dir /'prompts/course_objectives.txt')
        chain = prompt | self.model

        response = chain.invoke(
                        {
                            'grade' : self.grade,

                            'subject' : self.subject,

                            'Syllabus_type' : self.Syllabus_type,

                            "instructions" : self.instructions,

                            'course_description': course_description
                        }).content

        return self.Validator(response)

    def course_outline(self,course_description:str,course_objectives:str) -> str:
        # Structured outline explaining topics and subtopics covered in a timeline

        outline_prompt = self.build_prompt(current_dir / 'prompts/course_outline.txt')

        chain = outline_prompt | self.model
        response = chain.invoke(
                         {
                             'grade' : self.grade,
                             'subject' : self.subject,
                             'Syllabus_type' : self.Syllabus_type,
                             'instructions' : self.instructions,
                             'course_objectives' : course_objectives,
                             'course_description' : course_description,
                         }).content

        return self.Validator(response)

    def grading_policy(self,course_outline:str) -> str:
        # Assessment and marks policy and percentage distribtution

        prompt = self.build_prompt(current_dir / 'prompts/grading_policy.txt')
        chain = prompt | self.model

        response = chain.invoke(
                        {
                            'grade' : self.grade,
                            'subject' : self.subject,
                            'Syllabus_type' : self.Syllabus_type,
                            "instructions" : self.instructions,
                            'course_outline' : course_outline
                        }).content

        return self.Validator(response)

    def rules_policies(self,course_outline:str) -> str:
        # Classroom rules and etiquette

        prompt = self.build_prompt(current_dir / 'prompts/rules_policies.txt')
        chain = prompt | self.model

        response = chain.invoke(
                        {
                            'grade' : self.grade,
                            'subject' : self.subject,
                            'Syllabus_type' : self.Syllabus_type,
                            "instructions" : self.instructions,
                            'course_outline' : course_outline
                        }).content

        return self.Validator(response)

    def study_materials(self,course_outline:str) -> str:
        # Important study resources and their specific function

        prompt = self.build_prompt(current_dir / 'prompts/study_materials.txt')
        chain = prompt | self.model

        response = chain.invoke(
                        {
                            'grade' : self.grade,
                            'subject' : self.subject,
                            'Syllabus_type' : self.Syllabus_type,
                            "instructions" : self.instructions,
                            'course_outline' : course_outline
                        }).content
        return self.Validator(response)

    def run(self,verbose=False):
        # Final function to run all prompts and combine final syllabus

        #? time.sleep to avoid the quota increase error because if you use a free trial these a limit of requests per minute

        course_description = self.course_description()
        #time.sleep(10)
        course_objectives = self.course_objectives(course_description)
        #time.sleep(10)
        course_outline = self.course_outline(course_description,course_objectives)
        #time.sleep(10)
        grading_policy = self.grading_policy(course_outline)
        #time.sleep(10)
        rules_policies = self.rules_policies(course_outline)
        #time.sleep(10)
        study_materials = self.study_materials(course_outline)

        response = {
            'course_description':course_description,
            'course_objectives' :course_objectives,
            'study_materials'   :study_materials,
            'course_outline'    :course_outline,
            'grading_policy'    :grading_policy,
            'rules_policies'    :rules_policies
        }

        if verbose :
            for section_name,section_content in response.items():
                print("\n"+"*"*100)
                print(f"The results of {section_name} are :\n {section_content}")

        return response

# Class to generate PDF document with entire syllabus formatted with headings, tables, etc.
class PDFGenerator:
    def __init__(self, grade, subject):
        self.grade = grade
        self.subject = subject

    def remove_markdown(self, text):
        """Remove markdown syntax from text."""
        text = re.sub(r'#+ ', '', text)
        text = re.sub(r'\*{1,2}(.*?)\*{1,2}', r'\1', text)
        text = re.sub(r'^[*-]\s+', '', text, flags=re.MULTILINE)
        return text

    def group_dicts_by_subtopics(self,dicts,table_size=50):
        grouped_dicts = []
        current_group = []
        current_count = 0
        for item in dicts:
            subtopics_length = len(item['subtopics'])

            if current_count + subtopics_length > table_size:
                grouped_dicts.append(current_group)
                current_group = [item]
                current_count = subtopics_length
            else:
                current_group.append(item)
                current_count += subtopics_length

        if current_group:
            grouped_dicts.append(current_group)

        return grouped_dicts

    def split_into_chunks(self, text, max_length=101):
        words = text.split()
        chunks = []
        current_chunk = ""

        for word in words:
            if len(current_chunk) + len(word) + 1 > max_length:
                chunks.append(current_chunk)
                current_chunk = word
            else:
                if current_chunk:
                    current_chunk += " "
                current_chunk += word

        if current_chunk:
            chunks.append(current_chunk)

        return chunks

    def check_page_space(self,c, y_position,height, min_space=40):
        if y_position < min_space:
            c.showPage()
            c.setFont("Helvetica", 12)
            return height - 40
        return y_position

    def Wrap_line(self,List):
        lines = []
        for item in List:
            item_clean = f"• {self.remove_markdown(item)}"
            lines += self.split_into_chunks(item_clean, 100)

        return lines

    def generate_pdf(self,data: dict) -> BytesIO:
        buffer = BytesIO()
        c = canvas.Canvas(buffer, pagesize=letter)
        width, height = letter

        c.setFont("Helvetica-Bold", 36)
        title = f"Syllabus of {self.subject} "
        text_width = c.stringWidth(title, "Helvetica-Bold", 36)
        c.drawString((width - text_width) / 2, height / 2, title)

        c.showPage()
        y_position = height - 40
        c.setFont("Helvetica-Bold", 14)
        c.drawString(40, y_position, "Course Description")
        y_position -= 20
        c.setFont("Helvetica", 12)
        course_description = self.remove_markdown(data['course_description'])
        text_lines = self.split_into_chunks(course_description)

        for line in text_lines:
            c.drawString(40, y_position, line)
            y_position -= 12
            y_position = self.check_page_space(c, y_position,height)

        # Course Objectives
        y_position -= 20
        y_position = self.check_page_space(c, y_position,height)
        c.setFont("Helvetica-Bold", 14)
        c.drawString(40, y_position, "Course Objectives")
        y_position -= 20
        c.setFont("Helvetica", 12)
        Objectives_lines = self.Wrap_line(data['course_objectives'])

        for objective in Objectives_lines:
            c.drawString(40, y_position, objective)
            y_position -= 12
            y_position = self.check_page_space(c, y_position,height)

        # Grading Policy Table
        y_position -= 20
        c.setFont("Helvetica-Bold", 14)
        c.drawString(40, y_position, "Grading Policy")
        y_position -= 20
        grading_policy = data['grading_policy']
        table_data = [['Component', 'Coeff', 'Topics / Note']]
        table_data = table_data + [[item['Component'], f"{item['Coefficient']:.0%}", item['Note']] for item in grading_policy]

        grading_table = Table(table_data, colWidths=[80, 45, 435])
        grading_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
            ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, 0), 12),
            ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
            ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
            ('GRID', (0, 0), (-1, -1), 1, colors.black),
        ]))

        grading_table.wrapOn(c, 40, y_position)
        table_height = grading_table._height
        grading_table.drawOn(c, 40, y_position - table_height)
        y_position -= table_height + 30
        y_position = self.check_page_space(c, y_position,height)

        # Course Outline Table
        Groupes = self.group_dicts_by_subtopics(data['course_outline'])
        for group in Groupes:
            c.showPage()  # Start a new page for each table
            y_position = height - 40  # Reset y_position to the top of the new page
            c.setFont("Helvetica-Bold", 14)
            c.drawString(40, y_position, "Course Outline")
            y_position -= 20
            c.setFont("Helvetica", 10)
            table_data = [['Duration', 'Topic', 'Subtopics']]

            for item in group:
                subtopics = item.get('subtopics', [])
                splited_topic = self.split_into_chunks(item['topic'], 50)
                main_topic = '\n'.join(splited_topic)
                extra_newlines = max(0, len(subtopics) - len(splited_topic))
                main_topic += '\n' * extra_newlines

                subtopics_text = "\n".join(subtopics)
                week_info = item['duration']
                if len(subtopics) > 1:
                    week_info += '\n' * (len(subtopics) - 1)

                table_data.append([week_info, main_topic, subtopics_text])

            table = Table(table_data, colWidths=[50, 240, 270])
            table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, 0), 10),
                ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
                ('GRID', (0, 0), (-1, -1), 1, colors.black),
            ]))
            table.wrapOn(c, 40, y_position)
            table_height = table._height
            table.drawOn(c, 40, y_position - table_height)
            y_position -= table_height + 30


        # Study Materials
        c.showPage()
        y_position = height - 40
        c.setFont("Helvetica-Bold", 14)
        c.drawString(40, y_position, "Study Materials")
        y_position -= 20
        c.setFont("Helvetica", 12)
        table_data = [['Material','Purpose']]
        for material in data['study_materials']:
            material_clean = self.remove_markdown(material['material'])
            purpose_clean = self.remove_markdown(material['purpose'])
            material_clean = '\n'.join(self.split_into_chunks(material_clean, 56))
            purpose_clean = '\n'.join(self.split_into_chunks(purpose_clean, 56))
            table_data.append([material_clean, purpose_clean])
        table = Table(table_data, colWidths=[280,280])
        table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
            ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, 0), 10),
            ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
            ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
            ('GRID', (0, 0), (-1, -1), 1, colors.black),
        ]))
        table.wrapOn(c, 40, y_position)
        table_height = table._height
        table.drawOn(c, 40, y_position - table_height)
        y_position -= table_height + 30
        y_position = self.check_page_space(c, y_position,height)
        # Rules & Policies
        y_position -= 20
        y_position = self.check_page_space(c, y_position,height)
        c.setFont("Helvetica-Bold", 14)
        c.drawString(40, y_position, "Rules & Policies")
        y_position -= 20
        c.setFont("Helvetica", 12)
        rules_policies = data['rules_policies']
        for subtitle, content in rules_policies.items():
            y_position -= 10
            y_position = self.check_page_space(c, y_position,height)
            c.setFont("Helvetica-Bold", 12)
            c.drawString(40, y_position, subtitle)
            y_position -= 20
            c.setFont("Helvetica", 12)
            Rules_lines = self.Wrap_line(content)
            for rule in Rules_lines:
                c.drawString(40, y_position, rule)
                y_position -= 12
                y_position = self.check_page_space(c, y_position,height)

        y_position -= 20
        y_position = self.check_page_space(c, y_position, height)
        c.setFont("Helvetica-Bold", 14)
        c.drawString(40, y_position, "Memes")
        y_position -= 20
        c.setFont("Helvetica", 12)
        for meme_url in data.get('memes', []):
            link_text = f"• {meme_url}"
            c.drawString(40, y_position, link_text)
            c.linkURL(meme_url, (40, y_position - 10, 500, y_position + 5), relative=1)
            y_position -= 12
            y_position = self.check_page_space(c, y_position, height)


        c.save()
        buffer.seek(0)
        return buffer

# Similar functionality as PDF but for a word document
class WordGenerator:

    def __init__(self, grade, subject):
        self.grade = grade
        self.subject = subject
        self.document = Document()

    def remove_markdown(self, text):
        text = re.sub(r'#+ ', '', text)
        text = re.sub(r'\*{1,2}(.*?)\*{1,2}', r'\1', text)
        text = re.sub(r'^[*-]\s+', '', text, flags=re.MULTILINE)
        return text

    def add_title(self, title):
        self.document.add_heading(title, level=1)

    def add_heading(self, heading):
        self.document.add_heading(heading, level=2)

    def add_paragraph(self, text, bold=False, italic=False, underline=False):
        paragraph = self.document.add_paragraph()
        run = paragraph.add_run(text)
        run.bold = bold
        run.italic = italic
        run.underline = underline

    def add_table(self, data):
        table = self.document.add_table(rows=1, cols=len(data[0]))
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        for i, heading in enumerate(data[0]):
            hdr_cells[i].text = heading

        for row in data[1:]:
            row_cells = table.add_row().cells
            for i, value in enumerate(row):
                row_cells[i].text = value

    def generate_word(self, data: dict) -> BytesIO:
        title = f"Syllabus of {self.subject} {self.grade} Level"
        self.add_title(title)

        self.add_heading("Course Description")
        course_description = self.remove_markdown(data['course_description'])
        self.add_paragraph(course_description)

        self.add_heading("Course Objectives")
        for objective in data['course_objectives']:
            objective_clean = self.remove_markdown(objective)
            self.add_paragraph(f"• {objective_clean}")

        self.add_heading("Course Outline")
        table_data = [['Duration', 'Topic','Subtopics']]
        for item in data['course_outline']:
            main_topic = item['topic']
            subtopics = item.get('subtopics', [])
            if subtopics:
                subtopics = "\n".join([f"- {subtopic}" for subtopic in subtopics])
            table_data.append([item['duration'], main_topic, subtopics])

        self.add_table(table_data)


        self.add_heading("Grading Policy")
        grading_policy = data['grading_policy']
        table_data = [['Component', 'Coeff', 'Topics / Note']]
        for component in grading_policy:
            coefficient = f"{component['Coefficient']:.0%}"
            topics_note = self.remove_markdown(component['Note'])
            table_data.append([component['Component'], coefficient, topics_note])
        self.add_table(table_data)

        self.add_heading("Study Materials")
        table_data = [['Material', 'Purpose']]
        for material in data['study_materials']:
            material_clean = self.remove_markdown(material['material'])
            purpose_clean = self.remove_markdown(material['purpose'])
            table_data.append([material_clean, purpose_clean])
        self.add_table(table_data)

        self.add_heading("Rules & Policies")
        rules_policies = data['rules_policies']
        for subtitle, content in rules_policies.items():
            self.add_paragraph(subtitle, bold=True)
            for item in content:
                item_clean = self.remove_markdown(item)
                self.add_paragraph(f"• {item_clean}")

        self.add_heading("Memes")
        for meme_url in data.get('memes', []):
            self.add_paragraph(f"• {meme_url}")


        buffer = BytesIO()
        self.document.save(buffer)
        buffer.seek(0)
        return buffer